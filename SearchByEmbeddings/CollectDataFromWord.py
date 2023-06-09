from pandas import DataFrame
from docx import Document
import openai
import tiktoken
from setting import OPENAIKEY
openai.api_key=OPENAIKEY
import copy

filepath="data/城市建筑信息在城市作战中的应用.docx"
EMBEDDING_MODEL = "text-embedding-ada-002"
GPT_MODEL ="gpt-3.5-turbo"
MAX_TOKENS = 1600
BATCH_SIZE=1000

class CollectData():
    def __init__(self,filepath,outputpath):
        self.file=open(filepath,'rb')
        self.outputpath=outputpath
        self.document=Document(self.file)

    def halved_by_delimiter(self,string: str, delimiter: str = "\n") -> list[str, str]:
        """Split a string in two, on a delimiter, trying to balance tokens on each side."""
        chunks = string.split(delimiter)
        if len(chunks) == 1:
            return [string, ""]  # no delimiter found
        elif len(chunks) == 2:
            return chunks  # no need to search for halfway point
        else:
            total_tokens = self.num_tokens(string)
            halfway = total_tokens // 2
            best_diff = halfway
            for i, chunk in enumerate(chunks):
                left = delimiter.join(chunks[: i + 1])
                left_tokens = self.num_tokens(left)
                diff = abs(halfway - left_tokens)
                if diff >= best_diff:
                    break
                else:
                    best_diff = diff
            left = delimiter.join(chunks[:i])
            right = delimiter.join(chunks[i:])
            return [left, right]
    def truncated_string(self,
            string: str,
            model: str,
            max_tokens: int,
            print_warning: bool = True,
    ) -> str:
        """Truncate a string to a maximum number of tokens."""
        encoding = tiktoken.encoding_for_model(model)
        encoded_string = encoding.encode(string)
        truncated_string = encoding.decode(encoded_string[:max_tokens])
        if print_warning and len(encoded_string) > max_tokens:
            print(f"Warning: Truncated string from {len(encoded_string)} tokens to {max_tokens} tokens.")
        return truncated_string

    def split_strings_from_subsection(self,
            subsection: tuple[list[str], str],
            max_tokens: int = 1000,
            model: str = GPT_MODEL,
            max_recursion: int = 5,
    ) -> list[str]:
        """
        Split a subsection into a list of subsections, each with no more than max_tokens.
        Each subsection is a tuple of parent titles [H1, H2, ...] and text (str).
        """
        titles, text = subsection
        string = "\n\n".join(titles + [text])
        num_tokens_in_string = self.num_tokens(string)
        # if length is fine, return string
        if num_tokens_in_string <= max_tokens:
            return [string]
        # if recursion hasn't found a split after X iterations, just truncate
        elif max_recursion == 0:
            return [self.truncated_string(string, model=model, max_tokens=max_tokens)]
        # otherwise, split in half and recurse
        else:
            titles, text = subsection
            for delimiter in ["\n\n", "\n", ". "]:
                left, right = self.halved_by_delimiter(text, delimiter=delimiter)
                if left == "" or right == "":
                    # if either half is empty, retry with a more fine-grained delimiter
                    continue
                else:
                    # recurse on each half
                    results = []
                    for half in [left, right]:
                        half_subsection = (titles, half)
                        half_strings = self.split_strings_from_subsection(
                            half_subsection,
                            max_tokens=max_tokens,
                            model=model,
                            max_recursion=max_recursion - 1,
                        )
                        results.extend(half_strings)
                    return results
        # otherwise no split was found, so just truncate (should be very rare)
        return [self.truncated_string(string, model=model, max_tokens=max_tokens)]

    def num_tokens(self,text: str, model: str = GPT_MODEL) -> int:
        """Return the number of tokens in a string."""
        encoding = tiktoken.encoding_for_model("gpt2")
        return len(encoding.encode(text))

    def title(self):
        return self.document.core_properties.title

    def splittochunks(self):
        strings = []
        sections=self.getheadingwithparagraph()
        for section in sections:
            strings.extend(self.split_strings_from_subsection(section, max_tokens=MAX_TOKENS))

        print(f"{len(sections)} Wikipedia sections split into {len(strings)} strings.")
        print(strings)
        return strings

    def getheadingwithparagraph(self):
        res=[]
        level=1
        for i, para in enumerate(self.document.paragraphs):
            # Check if the paragraph has a heading style
            if para.style.name.startswith('Heading'):
                _level=para.style.name.split(" ")[1]
                # Get the text of the heading
                if int(_level)==1:
                    heading=[]
                    heading.append(para.text)
                else:
                    heading.append(para.text)
                # Get the text of the paragraph after the heading
                if i + 1 < len(self.document.paragraphs):
                    next_para = self.document.paragraphs[i + 1]
                    if next_para.style.name.startswith('Normal'):
                        next_para_text = next_para.text
                        #res.append({"title":self.title(),"heading":copy.deepcopy(heading),"content":next_para_text,"tokens":self.num_tokens(next_para_text)})
                        res.append((copy.deepcopy(heading),next_para_text))
        return res

    def embedchunks(self):
        embeddings = []
        strings=self.splittochunks()

        for batch_start in range(0, len(strings), BATCH_SIZE):
            batch_end = batch_start + BATCH_SIZE
            batch = strings[batch_start:batch_end]
            print(f"Batch {batch_start} to {batch_end - 1}")
            response = openai.Embedding.create(model=EMBEDDING_MODEL, input=batch)
            for i, be in enumerate(response["data"]):
                assert i == be["index"]  # double check embeddings are in same order as input
            batch_embeddings = [e["embedding"] for e in response["data"]]
            embeddings.extend(batch_embeddings)

        df = DataFrame({"text": strings, "embedding": embeddings})
        SAVE_PATH = self.outputpath
        df.to_csv(SAVE_PATH, index=False)
        return df

    def __del__(self):
        self.file.close()

